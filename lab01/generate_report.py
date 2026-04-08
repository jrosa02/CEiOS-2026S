"""Generowanie sprawozdania Word – Zadanie 3

Uruchomienie:
    uv run python lab01/generate_report.py
"""

from __future__ import annotations

import matplotlib
matplotlib.use("Agg")  # non-interactive backend – brak GUI

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from zadanie3 import (
    assign_zones,
    build_hourly_df,
    compute_costs,
    load_config,
)

LAB_DIR = Path(__file__).parent

# ---------------------------------------------------------------------------
# Stałe tekstowe
# ---------------------------------------------------------------------------

_FIXED_LABELS = {
    "sieciowa_stala": "Opłata sieciowa stała (OSD)",
    "abonamentowa":   "Opłata abonamentowa (OSD)",
    "mocowa":         "Opłata mocowa",
}

_TARIFF_ZONE_HOURS: dict[str, dict[str, str]] = {
    "G11": {
        "all_day": "wszystkie godziny doby (0–23), każdego dnia",
    },
    "G12": {
        "day":   "6:00–13:00 i 15:00–22:00 (każdego dnia tygodnia)",
        "night": "13:00–15:00 i 22:00–6:00 (każdego dnia tygodnia)",
    },
    "G12w": {
        "peak":     "6:00–13:00 i 15:00–22:00 (tylko dni robocze)",
        "off_peak": "13:00–15:00 i 22:00–6:00 (dni robocze) + wszystkie godziny sobót, niedziel i świąt",
    },
    "G13": {
        "afternoon_peak": "19:00–22:00 (tylko dni robocze) – godziny letnie",
        "morning_peak":   "7:00–13:00 (tylko dni robocze) – godziny letnie",
        "remaining":      "pozostałe godziny doby + wszystkie godziny sobót, niedziel i świąt",
    },
}

_TARIFF_NOTES: dict[str, str] = {
    "G11": (
        "Taryfa jednostrefowa – cena zakupu i dystrybucji jest identyczna "
        "o każdej porze dnia i tygodnia."
    ),
    "G12": (
        "Taryfa dwustrefowa obowiązuje 7 dni w tygodniu bez wyjątku – "
        "podział na dzień i noc jest stosowany również w soboty, niedziele i święta."
    ),
    "G12w": (
        "Taryfa dwustrefowa weekendowa różnicuje ceny w zależności od dnia tygodnia. "
        "Strefa szczytowa (droższa) obowiązuje wyłącznie w dni robocze. "
        "W soboty, niedziele oraz 13 dni ustawowo wolnych od pracy "
        "obowiązuje przez całą dobę tańsza strefa pozaszczytowa."
    ),
    "G13": (
        "Taryfa trzystrefowa z dwoma odrębnymi szczytami. "
        "Oba szczyty (popołudniowy i przedpołudniowy) obowiązują wyłącznie w dni robocze. "
        "W dni wolne od pracy (weekendy + 13 świąt) przez całą dobę "
        "obowiązuje najtańsza Strefa III. "
        "Uproszczenie zgodnie z instrukcją: godziny strefy letniej stosowane przez cały rok."
    ),
}


# ---------------------------------------------------------------------------
# Pomocnicze funkcje formatujące dokument
# ---------------------------------------------------------------------------

def _bold(run) -> None:
    run.bold = True


def _heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def _add_table_header(table, headers: list[str]) -> None:
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # szare tło nagłówka
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = tc.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "D9D9D9",
        })
        tcPr.append(shd)


def _add_row(table, values: list[str], right_cols: set[int] | None = None) -> None:
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = v
        align = WD_ALIGN_PARAGRAPH.RIGHT if (right_cols and i in right_cols) else WD_ALIGN_PARAGRAPH.LEFT
        cell.paragraphs[0].alignment = align


def _plnf(v: float, decimals: int = 2) -> str:
    fmt = f"{{:,.{decimals}f}}"
    return fmt.format(v).replace(",", "\u00a0")  # narrow non-breaking space jako separator tysięcy


# ---------------------------------------------------------------------------
# Sekcje dokumentu
# ---------------------------------------------------------------------------

def _write_intro(doc: Document, cfg: dict, df) -> None:
    import pandas as pd

    meta = cfg["metadata"]
    n_days = df["date"].nunique()
    n_work_free = int(df.groupby("date")["is_work_free"].first().sum())
    n_holidays = len(cfg["holidays_2026"]["dates"])
    n_weekends = n_work_free - n_holidays
    total_energy = df["energy_kwh"].sum()

    _heading(doc, "1. Dane wejściowe i założenia", level=1)

    p = doc.add_paragraph()
    p.add_run("Lokalizacja: ").bold = True
    p.add_run(f"{meta['location']}, rok {meta['year']}.")

    p = doc.add_paragraph()
    p.add_run("Sprzedawca: ").bold = True
    p.add_run("TAURON Sprzedaż sp. z o.o.")

    p = doc.add_paragraph()
    p.add_run("Podstawa prawna cen zakupu: ").bold = True
    p.add_run(
        "Taryfa zatwierdzona przez Prezesa URE, decyzja DRE.WRE.4211.53.10.2025.ASa1 "
        "z dnia 17 grudnia 2025 r., obowiązująca od 1 stycznia 2026 r."
    )

    p = doc.add_paragraph()
    p.add_run("Ceny dystrybucji: ").bold = True
    p.add_run(
        "wartości przybliżone na podstawie taryfy TAURON Dystrybucja S.A. "
        "(stawki netto, bez VAT, uwzględniające składnik zmienny sieciowy, "
        "stawkę jakościową, opłatę OZE i kogeneracyjną)."
    )

    p = doc.add_paragraph()
    p.add_run("Liczba dni w roku: ").bold = True
    p.add_run(f"{n_days} dni, {len(df)} godzin.")

    p = doc.add_paragraph()
    p.add_run("Dni wolne od pracy: ").bold = True
    p.add_run(
        f"{n_work_free} dni łącznie "
        f"({n_weekends} dni weekendowych + {n_holidays} dni ustawowo wolnych od pracy)."
    )
    p.add_run(
        " Dni wolne mają znaczenie dla taryf G12w i G13, "
        "w których przez całą dobę obowiązuje tańsza strefa pozaszczytowa / strefa III."
    )

    p = doc.add_paragraph()
    p.add_run("Roczne zużycie energii: ").bold = True
    p.add_run(
        f"{_plnf(total_energy)} kWh "
        "(profil godzinowy identyczny każdego dnia roku – tabela z instrukcji)."
    )

    p = doc.add_paragraph()
    p.add_run("Stawka VAT: ").bold = True
    p.add_run(f"{int(meta['vat_rate'] * 100)} %.")

    p = doc.add_paragraph()
    p.add_run("Cykl rozliczeniowy: ").bold = True
    p.add_run(
        f"co {meta['billing_period_months']} miesiące "
        f"({meta['billing_periods_per_year']} okresów rozliczeniowych w roku)."
    )

    # Tabela świąt
    _heading(doc, "1.1. Dni ustawowo wolne od pracy uwzględnione w obliczeniach", level=2)
    doc.add_paragraph(
        "Poniżej wymieniono 13 dni ustawowo wolnych od pracy w 2026 r. "
        "(art. 1 ustawy z dnia 18 stycznia 1951 r. o dniach wolnych od pracy). "
        "W tych dniach, podobnie jak w soboty i niedziele, taryfy G12w i G13 "
        "stosują ceny strefy pozaszczytowej / strefy III przez całą dobę."
    )

    holidays = cfg["holidays_2026"]["dates"]
    holiday_names = [
        "Nowy Rok", "Objawienie Pańskie (Trzech Króli)",
        "Niedziela Wielkanocna", "Poniedziałek Wielkanocny",
        "Święto Pracy", "Święto Konstytucji 3 Maja",
        "Zesłanie Ducha Świętego (Zielone Świątki)", "Boże Ciało",
        "Wniebowzięcie NMP", "Wszystkich Świętych",
        "Święto Niepodległości", "Boże Narodzenie (I dzień)", "Boże Narodzenie (II dzień)",
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    _add_table_header(tbl, ["Data", "Nazwa święta"])
    for date, name in zip(holidays, holiday_names):
        _add_row(tbl, [date, name])


def _write_tariff_section(
    doc: Document,
    section_no: int,
    tkey: str,
    tariff_cfg: dict,
    result: dict,
    vat: float,
    df_t,
) -> None:
    import pandas as pd

    zones = tariff_cfg["zones"]
    zone_hours = _TARIFF_ZONE_HOURS.get(tkey, {})

    _heading(doc, f"{section_no}. Taryfa {tkey} – {tariff_cfg['description']}", level=1)

    # Nota metodyczna
    note = _TARIFF_NOTES.get(tkey, "")
    if note:
        p = doc.add_paragraph(note)
        p.paragraph_format.space_after = Pt(6)

    # --- Definicja stref ---
    _heading(doc, f"{section_no}.1. Definicja stref czasowych", level=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    _add_table_header(tbl, ["Strefa", "Godziny obowiązywania", "Typ dnia"])
    for zk, zc in zones.items():
        hours_desc = zone_hours.get(zk, "—")
        if "all_hours" in zc:
            day_type = "wszystkie dni (pon–nie)"
        elif "workday_hours" in zc:
            day_type = "dni robocze; w dni wolne → strefa domyślna"
        else:
            day_type = "domyślna (pozostałe godziny i wszystkie dni wolne)"
        _add_row(tbl, [zc["name"], hours_desc, day_type])

    # Weryfikacja podziału dni
    _heading(doc, f"{section_no}.2. Podział godzin roku między strefy", level=2)
    n_work_free_days = int(df_t.groupby("date")["is_work_free"].first().sum())
    n_workdays = df_t["date"].nunique() - n_work_free_days
    doc.add_paragraph(
        f"W roku 2026 było {n_workdays} dni roboczych i {n_work_free_days} dni wolnych "
        f"(weekendy + święta). Poniższa tabela pokazuje, ile godzin i kWh "
        f"zostało przypisanych do każdej strefy."
    )

    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    _add_table_header(tbl2, ["Strefa", "Liczba godzin", "Udział [%]", "Energia [kWh]"])
    total_h = len(df_t)
    total_e = df_t["energy_kwh"].sum()
    for zr in result["zone_rows"]:
        hours_in_zone = int(round(zr["energy_kwh"] / (total_e / total_h)))
        pct = zr["energy_kwh"] / total_e * 100
        _add_row(tbl2, [
            zr["zone_name"],
            f"{hours_in_zone:,}".replace(",", "\u00a0"),
            f"{pct:.1f} %",
            _plnf(zr["energy_kwh"]),
        ], right_cols={1, 2, 3})

    # --- Ceny jednostkowe ---
    _heading(doc, f"{section_no}.3. Jednostkowe ceny zakupu i dystrybucji", level=2)
    doc.add_paragraph(
        "Wszystkie ceny są wartościami netto (bez VAT 23 %). "
        "Ceny zakupu zawierają podatek akcyzowy (0,005 PLN/kWh)."
    )
    tbl3 = doc.add_table(rows=1, cols=5)
    tbl3.style = "Table Grid"
    _add_table_header(tbl3, [
        "Strefa",
        "Zakup netto [PLN/kWh]",
        "Dystrybucja netto [PLN/kWh]",
        "Razem netto [PLN/kWh]",
        "Razem brutto [PLN/kWh]",
    ])
    for zr in result["zone_rows"]:
        total_unit = zr["purchase_unit"] + zr["distribution_unit"]
        _add_row(tbl3, [
            zr["zone_name"],
            f"{zr['purchase_unit']:.4f}",
            f"{zr['distribution_unit']:.4f}",
            f"{total_unit:.4f}",
            f"{total_unit * (1 + vat):.4f}",
        ], right_cols={1, 2, 3, 4})

    # --- Koszty zmienne strefowe ---
    _heading(doc, f"{section_no}.4. Koszty zmienne według stref", level=2)
    tbl4 = doc.add_table(rows=1, cols=7)
    tbl4.style = "Table Grid"
    _add_table_header(tbl4, [
        "Strefa",
        "Energia [kWh]",
        "Zakup netto [PLN]",
        "Dystrybucja netto [PLN]",
        "Razem zm. netto [PLN]",
        "VAT 23 % [PLN]",
        "Razem zm. brutto [PLN]",
    ])
    for zr in result["zone_rows"]:
        var_net = zr["purchase_net"] + zr["distribution_net"]
        _add_row(tbl4, [
            zr["zone_name"],
            _plnf(zr["energy_kwh"]),
            _plnf(zr["purchase_net"]),
            _plnf(zr["distribution_net"]),
            _plnf(var_net),
            _plnf(var_net * vat),
            _plnf(var_net * (1 + vat)),
        ], right_cols={1, 2, 3, 4, 5, 6})

    # --- Opłaty stałe ---
    _heading(doc, f"{section_no}.5. Opłaty stałe (roczne)", level=2)
    doc.add_paragraph(
        "Opłaty stałe są wspólne dla wszystkich taryf grupy G. "
        "Podane wartości miesięczne pomnożono przez 12 miesięcy."
    )
    tbl5 = doc.add_table(rows=1, cols=4)
    tbl5.style = "Table Grid"
    _add_table_header(tbl5, [
        "Składnik opłaty",
        "Stawka miesięczna [PLN]",
        "Roczna netto [PLN]",
        "Roczna brutto [PLN]",
    ])
    fixed_monthly = {k: v / 12 for k, v in result["fixed_annual"].items()}
    for k, v_annual in result["fixed_annual"].items():
        _add_row(tbl5, [
            _FIXED_LABELS.get(k, k),
            _plnf(fixed_monthly[k]),
            _plnf(v_annual),
            _plnf(v_annual * (1 + vat)),
        ], right_cols={1, 2, 3})
    # Suma opłat stałych
    total_fixed = result["total_fixed_net"]
    _add_row(tbl5, [
        "RAZEM opłaty stałe",
        _plnf(sum(fixed_monthly.values())),
        _plnf(total_fixed),
        _plnf(total_fixed * (1 + vat)),
    ], right_cols={1, 2, 3})
    # Pogrubienie wiersza sumy
    for cell in tbl5.rows[-1].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # --- Podsumowanie taryfy ---
    _heading(doc, f"{section_no}.6. Podsumowanie kosztów taryfy {tkey}", level=2)
    tbl6 = doc.add_table(rows=1, cols=3)
    tbl6.style = "Table Grid"
    _add_table_header(tbl6, ["Składnik kosztu", "Netto [PLN/rok]", "Brutto [PLN/rok]"])

    var_purchase = result["total_purchase_net"]
    var_distrib = result["total_distribution_net"]
    var_total = var_purchase + var_distrib
    fixed_total = result["total_fixed_net"]
    total_net = result["total_net"]
    total_gross = result["total_gross"]

    rows_summary = [
        ("Zakup energii (zmienny)", var_purchase),
        ("Dystrybucja zmienna", var_distrib),
        ("Razem koszty zmienne", var_total),
        ("Opłaty stałe", fixed_total),
        ("ŁĄCZNIE", total_net),
    ]
    for label, net_val in rows_summary:
        _add_row(tbl6, [label, _plnf(net_val), _plnf(net_val * (1 + vat))],
                 right_cols={1, 2})
    # Pogrubienie wiersza łącznie
    for cell in tbl6.rows[-1].cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True


def _write_summary(doc: Document, all_results: dict[str, dict], vat: float) -> None:
    section_no = len(all_results) + 2
    _heading(doc, f"{section_no}. Zestawienie porównawcze taryf", level=1)

    doc.add_paragraph(
        "Poniższa tabela zestawia całkowite roczne koszty energii elektrycznej "
        "dla każdej z porównywanych taryf. Wszystkie wartości dotyczą "
        "identycznego rocznego profilu zużycia energii."
    )

    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = "Table Grid"
    _add_table_header(tbl, [
        "Taryfa",
        "Zakup energii netto [PLN]",
        "Dystrybucja zm. netto [PLN]",
        "Opłaty stałe netto [PLN]",
        "Łącznie netto [PLN]",
        "VAT 23 % [PLN]",
        "Łącznie brutto [PLN]",
    ])
    for tkey, res in all_results.items():
        _add_row(tbl, [
            tkey,
            _plnf(res["total_purchase_net"]),
            _plnf(res["total_distribution_net"]),
            _plnf(res["total_fixed_net"]),
            _plnf(res["total_net"]),
            _plnf(res["total_net"] * vat),
            _plnf(res["total_gross"]),
        ], right_cols={1, 2, 3, 4, 5, 6})

    best = min(all_results, key=lambda t: all_results[t]["total_gross"])
    worst = max(all_results, key=lambda t: all_results[t]["total_gross"])
    diff = all_results[worst]["total_gross"] - all_results[best]["total_gross"]
    pct = diff / all_results[worst]["total_gross"] * 100

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Najtańsza taryfa: ").bold = True
    p.add_run(f"{best}  ({_plnf(all_results[best]['total_gross'])} PLN brutto/rok)")

    p = doc.add_paragraph()
    p.add_run("Najdroższa taryfa: ").bold = True
    p.add_run(f"{worst}  ({_plnf(all_results[worst]['total_gross'])} PLN brutto/rok)")

    p = doc.add_paragraph()
    p.add_run("Różnica między najtańszą a najdroższą taryfą: ").bold = True
    p.add_run(f"{_plnf(diff)} PLN/rok ({pct:.1f} % kosztu taryfy {worst}).")

    doc.add_paragraph(
        "Wybór taryfy G13 pozwala na największe oszczędności dzięki temu, "
        "że najtańsza Strefa III obejmuje wszystkie godziny dni wolnych "
        "oraz godziny poza szczytami w dni robocze. "
        "Taryfa G11, mimo najprostszej struktury, jest najdroższa – "
        "płaci się jednolitą, wyższą stawkę przez całą dobę."
    )


# ---------------------------------------------------------------------------
# Główna funkcja
# ---------------------------------------------------------------------------

def generate_report() -> Path:
    cfg = load_config()
    vat = cfg["metadata"]["vat_rate"]
    df = build_hourly_df(cfg)

    all_results: dict[str, dict] = {}
    df_with_zones: dict[str, dict] = {}
    for tkey, tcfg in cfg["tariffs"].items():
        if tkey == "common":
            continue
        df_t = assign_zones(df, tcfg)
        result = compute_costs(df_t, tcfg, cfg)
        all_results[tkey] = result
        df_with_zones[tkey] = (df_t, tcfg)

    doc = Document()

    # Styl czcionki bazowej
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Tytuł
    title = doc.add_heading("Zadanie 3 – Wybór najlepszej taryfy na zakup energii elektrycznej", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph("Porównanie taryf G11, G12, G12w, G13 – TAURON S.A., Kraków, rok 2026")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True

    doc.add_paragraph()

    # Sekcje
    _write_intro(doc, cfg, df)

    section_no = 2
    for tkey, (df_t, tcfg) in df_with_zones.items():
        doc.add_page_break()
        _write_tariff_section(doc, section_no, tkey, tcfg, all_results[tkey], vat, df_t)
        section_no += 1

    doc.add_page_break()
    _write_summary(doc, all_results, vat)

    out_path = LAB_DIR / "zadanie3_sprawozdanie.docx"
    doc.save(out_path)
    print(f"Sprawozdanie zapisane: {out_path}")
    return out_path


if __name__ == "__main__":
    generate_report()
